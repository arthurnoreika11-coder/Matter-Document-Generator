from __future__ import annotations

import re
import tempfile
from pathlib import Path

from docx import Document


def _safe_docx_filename(name: str) -> str:
    base_name = Path(name).name
    stem = Path(base_name).stem
    safe_stem = re.sub(r"[^A-Za-z0-9._-]+", "_", stem).strip("._-")

    if not safe_stem:
        safe_stem = "letter-before-action"

    return f"{safe_stem}.docx"


def create_lba(
    lba_name: str,
    recipient_name: str,
    legal_basis: str,
    demands: str,
    output_dir: str | Path | None = None,
) -> Path:
    doc = Document()

    doc.add_heading("Letter before Action", level=1)

    doc.add_paragraph(f"Dear {recipient_name},")
    doc.add_paragraph("Please accept this letter as a formal notice of my intention to take legal action.")
    doc.add_paragraph(f"Legal basis: {legal_basis}")
    doc.add_paragraph(f"Demands: {demands}")
    doc.add_paragraph("I request that you respond to this letter within 14 days of receipt. Failure to do so will result in further legal action.")

    target_dir = Path(output_dir) if output_dir is not None else Path(tempfile.mkdtemp(prefix="matter-docx-"))
    target_dir.mkdir(parents=True, exist_ok=True)
    docx_path = target_dir / _safe_docx_filename(lba_name)

    doc.save(docx_path)

    return docx_path
