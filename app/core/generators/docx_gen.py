from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document

from app.core.generators.doc_setup import _safe_docx_filename

def create_lba(
    lba_name: str,
    recipient_name: str,
    legal_basis: str,
    demands: str,
    output_dir: str | Path | None = None,
) -> Path:
    doc = Document()

    doc.add_heading("Letter Before Action", level=1)

    doc.add_paragraph(f"Dear {recipient_name},")
    doc.add_paragraph("Please accept this letter as a formal notice of my intention to take legal action.")
    doc.add_paragraph(f"Legal basis: {legal_basis}")
    doc.add_paragraph(f"Demands: {demands}")
    doc.add_paragraph("I request that you respond to this letter within 14 days of receipt. Failure to do so will result in further legal action.")

    target_dir = Path(output_dir) if output_dir is not None else Path(tempfile.mkdtemp(prefix="matter-docx-"))
    target_dir.mkdir(parents=True, exist_ok=True)
    docx_path = target_dir / _safe_docx_filename(lba_name)

    doc.save(docx_path)

    return docx_path
